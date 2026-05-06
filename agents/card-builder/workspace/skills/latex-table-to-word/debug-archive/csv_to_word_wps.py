import csv, re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import rcParams

csv_path = Path(r'J:\Desktop\table-1775988874328.csv')
out_path = Path(r'J:\Desktop\table-1775988874328_wps稳定版.docx')
img_dir = Path(r'C:\Users\Administrator\.openclaw\workspace\tmp_formula_imgs_wps')
img_dir.mkdir(exist_ok=True)

# 尽量稳妥地读取 CSV
last_err = None
rows = None
for enc in ('utf-8-sig', 'utf-8', 'gb18030', 'gbk'):
    try:
        with csv_path.open('r', encoding=enc, newline='') as f:
            rows = list(csv.reader(f))
        if rows:
            break
    except Exception as e:
        last_err = e
if not rows:
    raise RuntimeError('读取 CSV 失败: %r' % (last_err,))

headers = rows[0]
data_rows = rows[1:]
formula_cache = {}

# 数学字体配置
rcParams['mathtext.fontset'] = 'dejavuserif'
rcParams['font.family'] = ['Times New Roman', 'SimSun', 'DejaVu Serif']

def render_formula(formula: str) -> Path:
    formula = formula.strip()
    if formula in formula_cache:
        return formula_cache[formula]
    key = 'f%04d' % (len(formula_cache) + 1)
    out = img_dir / (key + '.png')

    fig = plt.figure(figsize=(0.01, 0.01))
    fig.patch.set_alpha(0)
    txt = fig.text(0, 0, formula, fontsize=14)
    fig.canvas.draw()
    bbox = txt.get_window_extent(renderer=fig.canvas.get_renderer()).expanded(1.08, 1.25)
    w, h = bbox.width / fig.dpi, bbox.height / fig.dpi
    plt.close(fig)

    fig = plt.figure(figsize=(max(w, 0.3), max(h, 0.3)))
    fig.patch.set_alpha(0)
    fig.text(0, 0, formula, fontsize=14)
    fig.savefig(str(out), dpi=300, transparent=True, bbox_inches='tight', pad_inches=0.03)
    plt.close(fig)
    formula_cache[formula] = out
    return out

def set_run_fonts(run, ascii_font='Times New Roman', east_font='宋体', size=10.5):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_font)
    run.font.size = Pt(size)


def add_mixed(paragraph, text: str):
    text = (text or '').replace('\r', ' ').replace('\n', ' ')
    parts = re.split(r'(\$.*?\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('$') and part.endswith('$'):
            latex = part.strip()
            if latex == '$$':
                continue
            img = render_formula(latex)
            run = paragraph.add_run()
            run.add_picture(str(img))
        else:
            run = paragraph.add_run(part)
            set_run_fonts(run)

doc = Document()

# 默认正文样式
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)

# 直接做表格，避免额外标题影响使用
cols = len(headers)
rows_n = len(data_rows) + 1
table = doc.add_table(rows=rows_n, cols=cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 表头
for j, h in enumerate(headers):
    cell = table.cell(0, j)
    p = cell.paragraphs[0]
    p.text = ''
    run = p.add_run(h)
    set_run_fonts(run, size=10.5)
    run.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# 内容
for i, row in enumerate(data_rows, start=1):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        p = cell.paragraphs[0]
        p.text = ''
        add_mixed(p, val)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# 粗略列宽
try:
    table.columns[0].width = Inches(0.9)
    if cols > 1:
        table.columns[1].width = Inches(3.0)
    if cols > 2:
        table.columns[2].width = Inches(4.0)
except Exception:
    pass

doc.save(str(out_path))
print(str(out_path))
