import csv, io, os, re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

csv_path = Path(r'J:\Desktop\table-1775988874328.csv')
out_path = Path(r'J:\Desktop\table-1775988874328.docx')
img_dir = Path(r'C:\Users\Administrator\.openclaw\workspace\tmp_formula_imgs')
img_dir.mkdir(exist_ok=True)

# Read CSV with UTF-8 BOM fallback
for enc in ('utf-8-sig', 'utf-8', 'gb18030'):
    try:
        with csv_path.open('r', encoding=enc, newline='') as f:
            rows = list(csv.reader(f))
        break
    except Exception:
        rows = None
if not rows:
    raise RuntimeError('无法读取 CSV 文件')

headers = rows[0]
data_rows = rows[1:]
formula_cache = {}

def render_formula(formula: str) -> Path:
    formula = formula.strip()
    if formula in formula_cache:
        return formula_cache[formula]
    key = 'f%04d' % (len(formula_cache) + 1)
    out = img_dir / (key + '.png')
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.patch.set_alpha(0)
    text = fig.text(0, 0, formula, fontsize=12)
    fig.canvas.draw()
    bbox = text.get_window_extent(renderer=fig.canvas.get_renderer()).expanded(1.05, 1.2)
    w, h = bbox.width / fig.dpi, bbox.height / fig.dpi
    plt.close(fig)

    fig = plt.figure(figsize=(max(w, 0.2), max(h, 0.2)))
    fig.patch.set_alpha(0)
    fig.text(0, 0, formula, fontsize=12)
    fig.savefig(str(out), dpi=300, transparent=True, bbox_inches='tight', pad_inches=0.02)
    plt.close(fig)
    formula_cache[formula] = out
    return out

def add_mixed(paragraph, text: str):
    text = (text or '').replace('\r', ' ').replace('\n', ' ')
    parts = re.split(r'(\$.*?\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('$') and part.endswith('$'):
            latex = part[1:-1].strip()
            if not latex:
                continue
            img = render_formula(f'${latex}$')
            run = paragraph.add_run()
            run.add_picture(str(img))
        else:
            paragraph.add_run(part)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)

doc.add_heading('CSV 转 Word（公式已格式化）', level=1)
table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

for j, h in enumerate(headers):
    cell = table.cell(0, j)
    cell.text = h
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

for i, row in enumerate(data_rows, start=1):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        add_mixed(p, val)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# simple width tuning by content type
try:
    table.columns[0].width = Inches(0.8)
    if len(headers) > 1:
        table.columns[1].width = Inches(2.8)
    if len(headers) > 2:
        table.columns[2].width = Inches(3.8)
except Exception:
    pass

doc.save(str(out_path))
print(str(out_path))
