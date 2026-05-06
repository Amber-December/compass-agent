import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter

input_path = Path(r'J:\Desktop\TEST.docx')
out_path = Path(r'J:\Desktop\TEST_重做公式对象版.docx')
xsl_path = Path(r'C:\Users\Administrator\.openclaw\workspace\skills\latex-table-to-word\assets\MML2OMML.XSL')

xslt = etree.XSLT(etree.parse(str(xsl_path)))

def latex_to_omml(latex: str):
    latex = latex.strip()
    if not latex:
        return None
    latex = latex.replace('∑', r'\sum ').replace('×', r'\times ').replace('−', '-')
    latex = latex.replace('β', r'\beta ').replace('ρ', r'\rho ')
    latex = re.sub(r'([A-Za-z]+)_([A-Za-z0-9]+)', r'\1_{\2}', latex)
    mathml = latex2mathml.converter.convert(latex)
    mml_root = etree.fromstring(mathml.encode('utf-8'))
    omml = xslt(mml_root)
    return etree.fromstring(etree.tostring(omml))


def append_text(paragraph, text):
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)


def append_omml(paragraph, omml_el):
    if omml_el is not None:
        paragraph._element.append(parse_xml(etree.tostring(omml_el)))


def mixed_append(paragraph, text):
    token_pat = r'([A-Za-z]+_[A-Za-z0-9]+|[A-Za-z]+\^\([^)]*\)|\([^)]*[A-Za-z]_[A-Za-z0-9]+[^)]*\))'
    parts = re.split(token_pat, text)
    for part in parts:
        if not part:
            continue
        if re.fullmatch(token_pat, part):
            try:
                append_omml(paragraph, latex_to_omml(part))
            except Exception:
                append_text(paragraph, part)
        else:
            append_text(paragraph, part)


def normalize_joined(s: str) -> str:
    s = s.replace('\u200b', '')
    s = re.sub(r'\s+', ' ', s).strip()
    s = re.sub(r'\b([A-Za-z]+) ([a-z]{1,12}|\d+|max)\b', r'\1_\2', s)
    s = s.replace(' _', '_')
    s = s.replace('_ ', '_')
    s = s.replace(' =', ' =').replace('( ', '(').replace(' )', ')')
    s = s.replace(' ，', '，').replace(' 。', '。').replace(' ：', '：')
    return s

src = Document(str(input_path))
raw_paras = [''.join(r.text for r in p.runs) for p in src.paragraphs]
raw_paras = [p for p in raw_paras if p.strip() and p.strip() != '\u200b']

# group fragmented paragraphs into logical lines
lines = []
buf = []
for p in raw_paras:
    t = p.replace('\u200b', '').strip()
    if not t:
        continue
    if re.match(r'^[SMC]-\d+\s', t) and buf:
        lines.append(normalize_joined(' '.join(buf)))
        buf = [t]
    else:
        buf.append(t)
if buf:
    lines.append(normalize_joined(' '.join(buf)))

out = Document()
style = out.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)

for line in lines:
    p = out.add_paragraph()
    m = re.match(r'^([SMC]-\d+\s+[^：]+：)(.*)$', line)
    if m:
        prefix, rest = m.group(1), m.group(2).strip()
        append_text(p, prefix)
        if '，其中' in rest:
            formula, tail = rest.split('，其中', 1)
            try:
                append_omml(p, latex_to_omml(formula))
            except Exception:
                append_text(p, formula)
            append_text(p, '，其中')
            mixed_append(p, tail)
        else:
            mixed_append(p, rest)
    else:
        mixed_append(p, line)

out.save(str(out_path))
print(out_path)
print('lines', len(lines))
for i,l in enumerate(lines[:8],1):
    print(i, l)
