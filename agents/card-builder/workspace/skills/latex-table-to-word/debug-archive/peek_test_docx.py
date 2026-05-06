from docx import Document
from pathlib import Path
p = Path(r'J:\Desktop\TEST.docx')
doc = Document(str(p))
for i,p in enumerate(doc.paragraphs[:40],1):
    txt=''.join(r.text for r in p.runs)
    if txt.strip():
        print(i, repr(txt[:300]))
for t_i, table in enumerate(doc.tables[:3],1):
    print('TABLE', t_i)
    for r_i,row in enumerate(table.rows[:5],1):
        vals=[]
        for cell in row.cells:
            vals.append('|'.join(p.text for p in cell.paragraphs))
        print(r_i, vals)
